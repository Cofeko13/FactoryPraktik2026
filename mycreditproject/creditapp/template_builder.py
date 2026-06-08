"""Создание .docx-шаблонов с плейсхолдерами Jinja2 для docxtpl."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _add_field_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (label, placeholder) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = placeholder


def build_anketa(path):
    doc = Document()
    _add_title(doc, 'АНКЕТА ЗАЕМЩИКА (физического лица)')
    doc.add_paragraph()
    _add_field_table(doc, [
        ('Фамилия, имя и отчество', '{{ full_name }}'),
        ('Дата рождения', '{{ birthdate }}'),
        ('Место рождения', '{{ birth_place }}'),
        ('Гражданство', '{{ citizenship }}'),
        ('Место работы и занимаемая должность', '{{ work_place }}'),
        ('Семейное положение', '{{ marital_status }}'),
        ('Документ, удостоверяющий личность', '{{ doc_type }}, {{ passport_full }}'),
        ('Адрес регистрации', '{{ registration_address }}'),
        ('ИНН', '{{ inn }}'),
        ('Телефон, факс', '{{ phone }}'),
        ('Зарплата (руб.)', '{{ salary }}'),
        ('Прочие доходы', '{{ other_income }}'),
        ('Обязательные платежи (руб.)', '{{ mandatory_payments }}'),
        ('Перечень собственности', '{{ property_list }}'),
        ('Сведения о кредитной истории', '{{ credit_history }}'),
        ('Цель кредита', '{{ loan_purpose }}'),
        ('Сумма и валюта', '{{ loan_amount_fmt }}'),
        ('Срок кредита (мес.)', '{{ loan_term }}'),
        ('Источники погашения', 'заработная плата, прочие доходы'),
        ('Предмет залога', '{{ collateral_subject }}'),
        ('Местонахождение залога', '{{ collateral_location }}'),
        ('Место хранения', '{{ collateral_storage }}'),
        ('Владелец залога', '{{ collateral_owner }}'),
        ('Сведения о поручителе(-ях)', '{{ guarantor_data }}'),
    ])
    doc.add_paragraph()
    doc.add_paragraph('Дата заполнения анкеты: {{ current_date }}')
    doc.add_paragraph('Подпись клиента: _________________')
    doc.add_paragraph('Подпись сотрудника, принявшего анкету: _________________')
    doc.save(path)


def build_soglasie_pd(path):
    doc = Document()
    _add_title(doc, 'СОГЛАСИЕ на обработку персональных данных')
    doc.add_paragraph()
    doc.add_paragraph(
        'Я, {{ full_name }} (далее – Клиент), даю согласие {{ bank_name }} '
        'на обработку моих персональных данных в соответствии с Федеральным законом '
        '№ 152-ФЗ «О персональных данных».'
    )
    doc.add_paragraph()
    _add_field_table(doc, [
        ('Ф.И.О.', '{{ full_name }}'),
        ('Дата и место рождения', '{{ birth_info }}'),
        ('Данные паспорта', '{{ passport_full }}'),
        ('Место регистрации', '{{ registration_address }}'),
        ('ИНН (при наличии)', '{{ inn }}'),
        ('СНИЛС (при наличии)', '{{ snils }}'),
        ('Фактическое место жительства', '{{ factual_address }}'),
        ('Телефон', '{{ phone }}'),
        ('E-mail', '{{ email }}'),
    ])
    doc.add_paragraph()
    doc.add_paragraph('Дата заполнения: {{ current_date }}')
    doc.add_paragraph('Подпись: _________________     {{ full_name }}')
    doc.save(path)


def build_soglasie_bki(path):
    doc = Document()
    _add_title(doc, 'СОГЛАСИЕ физического лица на получение информации, '
                     'характеризующей кредитную историю физического лица')
    doc.add_paragraph()
    doc.add_paragraph(
        'Я, {{ full_name }} (далее – Клиент), даю согласие {{ bank_name }} '
        'на получение информации, характеризующей мою кредитную историю, '
        'в соответствии с Федеральным законом № 218-ФЗ «О кредитных историях».'
    )
    doc.add_paragraph()
    _add_field_table(doc, [
        ('Ф.И.О.', '{{ full_name }}'),
        ('Дата и место рождения', '{{ birth_info }}'),
        ('Данные паспорта', '{{ passport_full }}'),
        ('Место регистрации', '{{ registration_address }}'),
        ('ИНН (при наличии)', '{{ inn }}'),
        ('СНИЛС (при наличии)', '{{ snils }}'),
        ('Фактическое место жительства', '{{ factual_address }}'),
        ('Код (пароль) Клиента (при наличии)', '{{ client_password }}'),
        ('Наименование БКИ', '{{ bki_name }}'),
    ])
    doc.add_paragraph()
    doc.add_paragraph('Дата заполнения: {{ current_date }}')
    doc.add_paragraph('Подпись: _________________     {{ full_name }}')
    doc.save(path)


def build_dogovor(path):
    doc = Document()
    _add_title(doc, 'КРЕДИТНЫЙ ДОГОВОР')
    doc.add_paragraph()
    doc.add_paragraph('г. ____________     {{ current_date_short }}')
    doc.add_paragraph()
    doc.add_paragraph(
        '{{ bank_name }}, именуемое в дальнейшем «Банк», в лице {{ bank_officer_position }} '
        '{{ bank_officer_name }}, действующего на основании {{ authority_basis }}, '
        'с одной стороны, и {{ full_name }}, именуемый(ая) в дальнейшем «Заёмщик», '
        'с другой стороны, заключили настоящий договор о нижеследующем:'
    )
    doc.add_paragraph()
    _add_field_table(doc, [
        ('Сумма кредита', '{{ loan_amount_fmt }}'),
        ('Цель кредита', '{{ loan_purpose }}'),
        ('Срок кредита (мес.)', '{{ loan_term }}'),
        ('Процентная ставка (% годовых)', '{{ interest_rate }}'),
        ('Полная стоимость кредита (%)', '{{ total_cost_percent }}'),
        ('Полная стоимость кредита (руб.)', '{{ total_cost_rub }}'),
        ('Переплата (руб.)', '{{ overpayment }}'),
        ('Неустойка (%)', '{{ penalty }}'),
        ('Предмет залога', '{{ collateral_subject }}'),
        ('Паспортные данные Заёмщика', '{{ passport_full }}'),
        ('Адрес регистрации', '{{ registration_address }}'),
        ('ИНН Заёмщика', '{{ inn }}'),
        ('Счёт клиента', '{{ client_account }}'),
        ('БИК банка', '{{ bank_bik }}'),
        ('К/с банка', '{{ correspondent_account }}'),
        ('ИНН банка', '{{ bank_inn }}'),
    ])
    doc.add_paragraph()
    doc.add_paragraph('Банк: _________________ / {{ bank_officer_name }} /')
    doc.add_paragraph('Заёмщик: _________________ / {{ full_name }} /')
    doc.save(path)


def build_zaklyuchenie(path):
    doc = Document()
    _add_title(doc, 'ЗАКЛЮЧЕНИЕ о возможности предоставления кредита')
    doc.add_paragraph()
    _add_field_table(doc, [
        ('ФИО заёмщика полностью', '{{ full_name }}'),
        ('Паспорт', '{{ passport_full }}'),
        ('Вид кредита', '{{ loan_type }}'),
        ('Сумма кредита', '{{ loan_amount_fmt }}'),
        ('Срок кредитования', '{{ loan_term }} мес.'),
        ('% за пользование кредитом', '{{ interest_rate }}'),
        ('Обеспечение', '{{ security }}'),
    ])
    doc.add_paragraph()
    doc.add_paragraph('Из представленных документов получены следующие сведения:')
    doc.add_paragraph()

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['Показатели', 'Значение', 'Балл']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for row_idx in range(1, 6):
        table.rows[row_idx].cells[0].text = f'{{{{ score_{row_idx}_indicator }}}}'
        table.rows[row_idx].cells[1].text = f'{{{{ score_{row_idx}_value }}}}'
        table.rows[row_idx].cells[2].text = f'{{{{ score_{row_idx}_score }}}}'

    table.rows[6].cells[0].text = 'Итоговый балл'
    table.rows[6].cells[1].text = ''
    table.rows[6].cells[2].text = '{{ total_score }}'

    doc.add_paragraph()
    _add_field_table(doc, [
        ('Среднемесячный чистый доход', '{{ net_income_calc }} руб.'),
        ('Платежеспособность заёмщика', '{{ solvency }} руб.'),
        ('Максимальная сумма кредита', '{{ max_loan_amount }} руб.'),
    ])
    doc.add_paragraph()
    doc.add_paragraph('{{ conclusion_text }}')
    doc.add_paragraph()
    doc.add_paragraph('Уполномоченный сотрудник банка')
    doc.add_paragraph('{{ current_date }}')
    doc.add_paragraph('_________________ / {{ bank_officer_name }} /')
    doc.save(path)


TEMPLATE_FILES = {
    'Анкета заемщика.docx': build_anketa,
    'Согласие.docx': build_soglasie_pd,
    'Согласие БКИ.docx': build_soglasie_bki,
    'Кредитный договор.docx': build_dogovor,
    'Заключение о возможности предоставления кредита.docx': build_zaklyuchenie,
}


def ensure_templates(templates_dir):
    templates_dir = Path(templates_dir)
    templates_dir.mkdir(parents=True, exist_ok=True)
    for filename, builder in TEMPLATE_FILES.items():
        path = templates_dir / filename
        if not path.exists():
            builder(str(path))
